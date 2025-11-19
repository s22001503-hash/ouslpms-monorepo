"""
Enhanced Document Extractor with Chunking Support
=================================================

Extracts text from PDF, DOCX, and images with intelligent chunking
for long documents.

Author: OUSL Print Management System  
Date: November 1, 2025
"""

import os
from typing import Optional, List, Dict
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract

import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from app.utils.text_chunker import TextChunker, Chunk


class DocumentExtractor:
    """Extract text from various document formats with chunking support"""
    
    def __init__(self):
        # Set Tesseract path for Windows
        if os.name == 'nt':
            pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        
        # Initialize chunker
        self.chunker = TextChunker(
            chunk_size=1500,      # words per chunk
            chunk_overlap=200,    # words overlap
            min_chunk_size=100    # minimum chunk size
        )
    
    def extract_text(
        self,
        file_path: str,
        enable_chunking: bool = True,
        chunk_strategy: str = "auto"
    ) -> Dict:
        """
        Extract text from document with optional chunking
        
        Args:
            file_path: Path to the document file
            enable_chunking: Whether to chunk long documents
            chunk_strategy: "auto", "fixed", or "semantic"
        
        Returns:
            dict with text content, chunks, and metadata
        """
        file_extension = Path(file_path).suffix.lower()
        
        # Extract raw text based on file type
        if file_extension == '.pdf':
            extraction_result = self._extract_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            extraction_result = self._extract_from_docx(file_path)
        elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            extraction_result = self._extract_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        full_text = extraction_result["text"]
        
        # Apply chunking if enabled
        if enable_chunking:
            chunks = self.chunker.chunk_text(full_text, strategy=chunk_strategy)
        else:
            # Create single chunk for entire document
            chunks = [Chunk(
                text=full_text,
                chunk_id=0,
                total_chunks=1,
                page_range=(1, extraction_result.get("page_count", 1)),
                word_count=len(full_text.split()),
                char_count=len(full_text),
                section_title="Full Document"
            )]
        
        # Prepare result
        result = {
            **extraction_result,
            "chunks": chunks,
            "is_chunked": len(chunks) > 1,
            "total_chunks": len(chunks),
            "chunk_strategy": chunk_strategy if len(chunks) > 1 else "none"
        }
        
        return result
    
    def _extract_from_pdf(self, file_path: str) -> Dict:
        """Extract text from PDF file"""
        text = ""
        page_count = 0
        
        try:
            # Method 1: Try pdfplumber (best for most PDFs)
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If no text extracted, try OCR (scanned PDF)
            if not text.strip():
                print("→ PDF appears to be scanned, using OCR...")
                text = self._ocr_pdf(file_path)
        
        except Exception as e:
            # Method 2: Fallback to PyPDF2
            print(f"→ pdfplumber failed ({e}), trying PyPDF2...")
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    page_count = len(pdf_reader.pages)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                
                # If still no text, use OCR
                if not text.strip():
                    text = self._ocr_pdf(file_path)
            
            except Exception as e2:
                print(f"→ PyPDF2 also failed ({e2}), using OCR...")
                text = self._ocr_pdf(file_path)
        
        return {
            "text": text.strip(),
            "page_count": page_count,
            "file_type": "pdf",
            "extraction_method": "text" if text else "ocr"
        }
    
    def _extract_from_docx(self, file_path: str) -> Dict:
        """Extract text from DOCX file"""
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += "\n" + row_text
        
        return {
            "text": text.strip(),
            "page_count": len(doc.sections),
            "file_type": "docx",
            "extraction_method": "text"
        }
    
    def _extract_from_image(self, file_path: str) -> Dict:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            return {
                "text": text.strip(),
                "page_count": 1,
                "file_type": "image",
                "extraction_method": "ocr"
            }
        
        except Exception as e:
            raise ValueError(f"Error extracting text from image: {e}")
    
    def _ocr_pdf(self, file_path: str) -> str:
        """Extract text from scanned PDF using OCR"""
        try:
            from pdf2image import convert_from_path
            
            # Convert PDF to images
            images = convert_from_path(file_path, dpi=300)
            
            # Extract text from each page
            text = ""
            for i, image in enumerate(images):
                print(f"  → OCR processing page {i+1}/{len(images)}...")
                page_text = pytesseract.image_to_string(image)
                text += page_text + "\n"
            
            return text
        
        except Exception as e:
            raise ValueError(f"Error performing OCR on PDF: {e}")


# Example usage and testing
if __name__ == "__main__":
    extractor = DocumentExtractor()
    
    # Test files (update paths as needed)
    test_files = {
        "short": "sample_short.pdf",      # < 10 pages
        "medium": "sample_medium.pdf",    # 10-30 pages
        "long": "sample_long.pdf"         # 100+ pages
    }
    
    for doc_type, file_path in test_files.items():
        if os.path.exists(file_path):
            print(f"\n{'='*60}")
            print(f"Testing {doc_type.upper()} document: {file_path}")
            print('='*60)
            
            result = extractor.extract_text(file_path, enable_chunking=True)
            
            print(f"File Type: {result['file_type']}")
            print(f"Page Count: {result['page_count']}")
            print(f"Extraction Method: {result['extraction_method']}")
            print(f"Is Chunked: {result['is_chunked']}")
            print(f"Total Chunks: {result['total_chunks']}")
            print(f"Chunk Strategy: {result['chunk_strategy']}")
            
            if result['is_chunked']:
                print(f"\nChunk Details:")
                for chunk in result['chunks'][:3]:  # Show first 3 chunks
                    print(f"\n  Chunk {chunk.chunk_id + 1}/{chunk.total_chunks}:")
                    print(f"    Section: {chunk.section_title}")
                    print(f"    Pages: {chunk.page_range[0]}-{chunk.page_range[1]}")
                    print(f"    Words: {chunk.word_count}")
                    print(f"    Preview: {chunk.text[:100]}...")
            else:
                print(f"\nFull Document:")
                print(f"  Words: {len(result['text'].split())}")
                print(f"  Preview: {result['text'][:200]}...")
